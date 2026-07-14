"""Rebuild the evaluation section of ``plank-road.docx`` from measured results.

The script preserves the source document and writes a separate revised DOCX.
It intentionally reports single-run and case-study limitations instead of
creating uncertainty estimates that are absent from the archived results.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "plank-road.docx"
OUTPUT = ROOT / "plank-road-evaluation-revised.docx"

RAINY_TRADEOFF = (
    ROOT
    / "results/experiments/suwon5a_weather_rainy/figures/"
    "fig2_accuracy_retraining_time_tradeoff.png"
)
SNOWY_TRADEOFF = (
    ROOT
    / "results/experiments/suwon5a_weather/figures/"
    "fig2_accuracy_retraining_time_tradeoff.png"
)
TAIL_TRAINING = (
    ROOT
    / "results/tail_training_motivation/plots/"
    "freeze_vs_split_cached_vs_rebuild_by_position.png"
)
DRIFT_VALIDITY = (
    ROOT
    / "results/drift_detection_validity/suwon5a_real_weather_scene_test/plots/"
    "exp1_timeseries_real_weather.png"
)
PRIVACY_MATRIX = (
    ROOT
    / "results/privacy_reconstruction_scene_split_paper_figures/"
    "privacy_reconstruction_attacks_paper_matrix.png"
)


def _find_unique_paragraph(document: Document, text: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected exactly one paragraph {text!r}, found {len(matches)}")
    return matches[0]


def _remove_between(start, end) -> None:
    node = start._p.getnext()
    while node is not None and node is not end._p:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node
    if node is None:
        raise RuntimeError("End paragraph was not found after the start paragraph")


def _move_before(element, anchor) -> None:
    anchor._p.addprevious(element)


def _set_run_font(run, size: float | None = None, *, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _add_paragraph(document: Document, anchor, text: str, style: str = "Body Text"):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    _move_before(paragraph._p, anchor)
    return paragraph


def _add_caption(document: Document, anchor, text: str):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    _set_run_font(run, 7.5)
    _move_before(paragraph._p, anchor)
    return paragraph


def _add_picture(document: Document, anchor, path: Path, caption: str):
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(3.36))
    _move_before(paragraph._p, anchor)
    return _add_caption(document, anchor, caption)


def _set_cell_margins(cell, top=40, start=50, bottom=40, end=50) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.first_child_found_in("w:tblBorders")
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for name, val, size, color in (
        ("top", "single", "8", "555555"),
        ("left", "nil", "0", "FFFFFF"),
        ("bottom", "single", "8", "555555"),
        ("right", "nil", "0", "FFFFFF"),
        ("insideH", "single", "4", "BFBFBF"),
        ("insideV", "nil", "0", "FFFFFF"),
    ):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_main_results_table(document: Document, anchor) -> None:
    caption = document.add_paragraph(style="Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run("表XX　真实天气场景下的端到端效果与开销（每个场景 1 次正式运行）")
    _set_run_font(run, 7.5)
    _move_before(caption._p, anchor)

    headers = ("场景", "方法", "F1↑", "训练(s)↓", "延迟(ms)↓")
    rows = (
        ("雨天", "Plank-road", "0.702", "71.13", "206.61"),
        ("雨天", "SURGEON", "0.474", "769.33", "342.93"),
        ("雨天", "CATR", "0.648", "262.86", "327.20"),
        ("雨天", "Ekya", "0.680", "206.26", "247.64"),
        ("雪天", "Plank-road", "0.693", "69.12", "81.65"),
        ("雪天", "SURGEON", "0.621", "350.82", "412.69"),
        ("雪天", "CATR", "0.623", "176.10", "85.93"),
        ("雪天", "Ekya", "0.638", "134.77", "243.71"),
    )
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Inches(0.43), Inches(0.86), Inches(0.48), Inches(0.78), Inches(0.81))
    for column, width in zip(table.columns, widths):
        column.width = width
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.w = width
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(round(sum(width.pt for width in widths) * 20)))
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for col, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[col]
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell)
        _shade_cell(cell, "D9EAF7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        _set_run_font(run, 6.8, bold=True)

    for values in rows:
        cells = table.add_row().cells
        for col, (value, width) in enumerate(zip(values, widths)):
            cell = cells[col]
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if value == "Plank-road":
                _shade_cell(cell, "EDF5FB")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            _set_run_font(run, 6.7, bold=(value == "Plank-road"))

    _set_table_borders(table)
    _move_before(table._tbl, anchor)
    _add_caption(
        document,
        anchor,
        "注：F1 为教师监督 F1；教师回放耗时不计入在线推理、传输与重训练开销。",
    )


def build() -> Path:
    document = Document(SOURCE)
    evaluation = _find_unique_paragraph(document, "评估")
    discussion = _find_unique_paragraph(document, "讨论")
    _remove_between(evaluation, discussion)

    _add_paragraph(document, discussion, "实验设置与评估协议", "Heading 2")
    _add_paragraph(
        document,
        discussion,
        "实验平台由一台配备 NVIDIA RTX A6000 GPU 和 96 GB 系统内存的云端服务器，以及 ThinkPad T14p、树莓派 4B（2 GB 内存）和 Jetson Orin NX（8 GB 内存）等边缘设备组成。系统支持 CNN 架构的 YOLO26n 与 Transformer 架构的 RF-DETR Nano 作为边缘检测模型，云端使用 RT-DETR-X 生成教师伪标签。端到端基线实验统一采用 RF-DETR Nano 作为学生模型；切分与隐私辅助实验进一步覆盖 YOLO26n 和 TinyNeXt-S。视频流来自 TSBOW 数据集的 Suwon#5a 交通监控场景，主对比使用雨天与雪天视频，以覆盖能见度下降、遮挡与天气扰动造成的数据分布变化。",
    )
    _add_paragraph(
        document,
        discussion,
        "为保证比较一致，四种方法使用相同的视频回放、学生模型与教师模型。检测效果采用教师监督 F1：教师与学生预测的置信度阈值均设为 0.6，类别感知匹配的 IoU 阈值设为 0.5。该指标衡量学生预测与教师伪标签的一致性，而不是人工标注意义上的真实准确率。教师回放仅用于离线评估，其耗时不计入在线推理、上传、标注或训练时延。系统开销报告平均在线推理时延、平均训练时间、平均上传字节数和原始帧暴露比例；缺失的 mAP 不由 F1 推算。归档的雨天和雪天主实验均为单边缘、单次正式运行，因此下文只报告观测值与相对差异，不给出方差或统计显著性结论。",
    )
    _add_paragraph(
        document,
        discussion,
        "对比方法包括三类代表性策略：（1）SURGEON，在边缘端冻结前缀并执行无监督测试时自适应，不向云端上传训练数据；（2）CATR，由云端教师标注上传帧，并依据教师监督 F1 的下降触发冻结式云端重训练；（3）Ekya，将帧上传至云端，在微型性能剖析后统一调度推理与持续学习。SURGEON 的上传量和原始帧暴露比例为结构性零值，不能解释为与云端方法具有相同训练路径。",
    )

    _add_paragraph(document, discussion, "整体检测效果与系统效率", "Heading 2")
    _add_paragraph(
        document,
        discussion,
        "表XX汇总了端到端主实验。在雨天场景中，Plank-road 的教师监督 F1 为 0.702，分别比 Ekya、CATR 和 SURGEON 高 0.021、0.054 和 0.228。其平均训练时间为 71.13 s，相比三种基线分别减少 65.5%、72.9% 和 90.8%；平均在线推理时延为 206.61 ms，分别降低 16.6%、36.9% 和 39.8%。因此，在该次雨天运行中，Plank-road 位于更高 F1 和更低训练时间构成的优势区域。",
    )
    _add_paragraph(
        document,
        discussion,
        "雪天场景呈现相同趋势。Plank-road 的教师监督 F1 为 0.693，比 Ekya、CATR 和 SURGEON 分别高 0.055、0.071 和 0.072；平均训练时间为 69.12 s，分别减少 48.7%、60.7% 和 80.3%。其平均在线推理时延为 81.65 ms，是四种方法中的最低值，相比 CATR、Ekya 和 SURGEON 分别降低 5.0%、66.5% 和 80.2%。这些结果表明，在两个已完成的真实天气运行中，切分尾部训练能够以较短的更新周期维持较高的教师一致性。",
    )
    _add_main_results_table(document, discussion)
    _add_picture(
        document,
        discussion,
        RAINY_TRADEOFF,
        "图XX　雨天场景中的教师监督 F1—平均训练时间权衡。每个点对应一次正式运行，未绘制方差椭圆。",
    )
    _add_picture(
        document,
        discussion,
        SNOWY_TRADEOFF,
        "图XX　雪天场景中的教师监督 F1—平均训练时间权衡。每个点对应一次正式运行，未绘制方差椭圆。",
    )
    _add_paragraph(
        document,
        discussion,
        "通信与数据暴露结果进一步刻画了这种权衡。雨天运行中，Plank-road 的平均上传量为 6.80 MB，低于 CATR 的 13.10 MB 和 Ekya 的 14.73 MB；雪天运行中相应数值为 3.83、4.01 和 10.27 MB。Plank-road 的原始帧暴露比例在雨天和雪天分别为 0.933 和 0.889，低于两个云端基线的 1.0，但仍接近完全暴露。这说明系统减少了部分原始帧传输，却没有消除教师标注所需的原始数据；隐私收益应与后续重建攻击实验结合解释。",
    )

    _add_paragraph(document, discussion, "动态适应与资源触发行为", "Heading 2")
    _add_paragraph(
        document,
        discussion,
        "在雨天运行中，Plank-road 记录了 14 次触发决策、17 个训练作业和 12 次已应用模型更新；雪天运行中相应为 8、13 和 6。用于绘图的事件配对规则仅保留能够与后续模型更新对应的触发，其中两个场景均有 2 次未配对触发被明确报告而未绘制。相比之下，CATR 在雨天/雪天分别应用 4/2 次更新，Ekya 为 1/1 次，SURGEON 为 1/1 次。该结果显示 Plank-road 采用了更细粒度的尾部更新，并能够延后或取消部分尚未形成更新闭环的触发；但现有结果没有包含去除 Lyapunov 队列的端到端消融，因此不能仅凭这些计数将全部收益归因于资源队列控制。",
    )

    _add_paragraph(document, discussion, "切分尾部训练消融", "Heading 2")
    _add_paragraph(
        document,
        discussion,
        "为隔离切分尾部训练的作用，我们在 RF-DETR Nano 上使用 512 个样本、10 个训练轮次、批大小 32，并在前部 25%、中部 50% 和后部 75% 三个切分位置各重复 5 次。实验比较原生冻结训练、TorchLens 冻结训练、切分后重建一次边界特征（split rebuild）以及直接复用已缓存边界特征（split cached）；所有模式均训练同一后缀参数集合，并以教师伪标签上的代理 mAP@[0.5:0.95] 评价训练前后变化。",
    )
    _add_paragraph(
        document,
        discussion,
        "缓存切分将三个位置的平均尾部训练时间从 40.64、32.73 和 31.93 s 降至 22.86、10.63 和 9.67 s，分别减少 43.8%、67.5% 和 69.7%。在中部切分处，代理 mAP 增量为 0.122，与 TorchLens 冻结训练的 0.124 基本接近；后部切分的增量为 0.066，略高于冻结训练的 0.060；前部切分的增量则由 0.162 降至 0.137。若缓存缺失，单次边界特征重建额外耗时 1.77–2.20 s，但三处的总更新时间仍比对应冻结训练低 29.6%–43.5%。结果说明缓存边界特征是主要加速来源，同时切分位置决定了可训练后缀规模与适应收益：切分越靠后，训练越快，但可获得的精度增量总体越小。",
    )
    _add_picture(
        document,
        discussion,
        TAIL_TRAINING,
        "图XX　不同切分位置下的尾部训练时间与代理检测质量。箱线图基于每种设置 5 次重复。",
    )

    _add_paragraph(document, discussion, "漂移信号与在线触发有效性", "Heading 2")
    _add_paragraph(
        document,
        discussion,
        "漂移有效性实验从晴天、雨天和雪天视频中各均匀采样 24 帧，并使用同帧 RT-DETR-X 伪标签进行离线评价。学生—教师微平均 F1 从晴天的 0.795 降至雨天的 0.562 和雪天的 0.497，表明恶劣天气造成了可观测的检测一致性退化。在 9 个窗口上，完整无标签漂移分数与 F1 下降的 Pearson 和 Spearman 相关系数分别为 0.948 和 0.812；输出熵、边界特征偏移及其指数滑动统计在该组窗口上的 ROC-AUC 和 PR-AUC 均为 1.0。在线回放中，完整 Plank-road 信号检测到唯一的有害漂移事件，未出现误触发，并在配置的容忍窗口内实现 0 帧延迟；仅使用置信度的基线虽然召回该事件，但产生 1 次误触发，触发 F1 为 0.667。",
    )
    _add_picture(
        document,
        discussion,
        DRIFT_VALIDITY,
        "图XX　真实天气流中无标签漂移分数与教师伪标签 F1 下降的同步变化。教师信息仅用于离线判定，不参与在线触发。",
    )
    _add_paragraph(
        document,
        discussion,
        "上述结果支持“输出不确定性 + 边界表征偏移”可作为有害天气漂移的在线代理，但其证据范围有限：当前仅包含一条按晴—雨—雪排列的序列、72 帧和 9 个窗口，且只有 1 个有害漂移事件。AUC=1.0 因而应视为受控回放中的可分性结果，而不能直接外推到更多摄像头、场景顺序或长期运行。",
    )

    _add_paragraph(document, discussion, "切分规划与隐私泄露量化", "Heading 2")
    _add_paragraph(
        document,
        discussion,
        "Graph Split Planner 的导出记录覆盖 RF-DETR Nano、YOLO26n 和 TinyNeXt-S 的 570、323 和 359 个候选边界。所有导出候选均通过运行时验证，具有可训练尾部且边界重放成功率为 1.0。候选中间特征相对输入张量的负载范围具有明显模型差异：RF-DETR Nano 为 0.056–1.313，TinyNeXt-S 为 0.025–1.000，而 YOLO26n 为 1.000–9.214。该结果说明中间特征并不天然小于原始输入，尤其 YOLO26n 的候选边界可能放大传输量，因此必须在隐私、可训练性和通信约束下按模型选择边界，而不能使用固定层号。",
    )
    _add_paragraph(
        document,
        discussion,
        "我们进一步在 first-compute 边界和目标隐私泄露分数 0.75、0.50、0.25 处执行 DRAG 与白盒特征反演。白盒攻击下，三个模型在 0.75、0.50 和 0.25 三个受约束切分点的重建目标 Object F1 均为 0；作为对照，TinyNeXt-S 的 first-compute 特征可被近乎完整恢复（SSIM=0.998、Object F1=1.0、实际泄露度=0.999），而目标分数 0.25 处降至 SSIM=0.066、Object F1=0、实际泄露度=0.072。DRAG 攻击下，RF-DETR Nano 在 first-compute 和 0.75 处的 Object F1 分别为 0.286 和 0.250，在 0.50 与 0.25 处降为 0；YOLO26n 与 TinyNeXt-S 的 DRAG Object F1 在本次样本上均为 0。定性重建也显示，较深边界通常更难保留可识别的车辆与道路语义。",
    )
    _add_picture(
        document,
        discussion,
        PRIVACY_MATRIX,
        "图XX　三个边缘模型在 DRAG 与白盒反演攻击下的重建案例；列标题为 first-compute 或目标隐私泄露分数。",
    )
    _add_paragraph(
        document,
        discussion,
        "隐私实验同时暴露了当前评估的边界：每个模型、攻击方法和切分点仅包含 1 个重建样本，且不同模型的实际泄露度并非严格单调。例如 YOLO26n 的白盒实际泄露度在各受约束边界间为 0.059–0.085。因而现有结果只能作为攻击压力测试和机制案例，支持在规划器中保留模型相关的隐私约束，但不足以证明参数比例代理在所有模型和场景上与真实重建风险严格一致。",
    )

    _add_paragraph(document, discussion, "评估结论与证据边界", "Heading 2")
    _add_paragraph(
        document,
        discussion,
        "综合现有结果，Plank-road 在雨天和雪天的单次端到端运行中同时获得了最高教师监督 F1、最短平均训练时间和较低在线推理时延；五次重复的尾部训练消融进一步表明，缓存边界特征能够在大体保持代理 mAP 增量的同时缩短训练时间。漂移回放和重建攻击则分别验证了无标签漂移信号与有害性能下降的对应关系，以及深层特征降低可恢复语义的可行性。当前证据尚未覆盖主对比的多次重复、多摄像头并发、不同带宽/云端负载下的 Lyapunov 参数消融，也未形成大样本隐私攻击统计；这些结果应在后续实验中补充后再用于显著性、可扩展性或普遍隐私保证的结论。",
    )

    document.core_properties.modified = datetime.now()
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
